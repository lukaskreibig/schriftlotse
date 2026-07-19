from __future__ import annotations

import html
import json
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from schriftlotse import __version__
from schriftlotse.domain import DocumentResult, PageResult
from schriftlotse.pagexml import NS


def original_text(result: DocumentResult) -> str:
    pages = []
    for page in result.pages:
        pages.append("\n".join(line.text for line in page.lines))
    return "\n\n--- Seite ---\n\n".join(pages).strip() + "\n"


def reading_text(result: DocumentResult) -> str:
    text = original_text(result)
    text = re.sub(r"(?<=\w)-\n(?=[a-zäöüß])", "", text)
    text = re.sub(r"(?<!\n)\n(?!\n|--- Seite ---)", " ", text)
    return re.sub(r"[ \t]+", " ", text).strip() + "\n"


def _polygon_points(points: list[tuple[int, int]]) -> str:
    return " ".join(f"{x},{y}" for x, y in points)


def _page_xml(
    result: DocumentResult,
    page: PageResult,
    destination: Path,
    image_filename: str | None = None,
) -> None:
    root = ET.Element(f"{{{NS}}}PcGts")
    metadata = ET.SubElement(root, f"{{{NS}}}Metadata")
    ET.SubElement(metadata, f"{{{NS}}}Creator").text = f"SchriftLotse {__version__}"
    page_node = ET.SubElement(
        root,
        f"{{{NS}}}Page",
        imageFilename=image_filename or page.source_path.name,
        imageWidth=str(page.width),
        imageHeight=str(page.height),
    )
    region_nodes: dict[str, ET.Element] = {}
    for item in page.regions:
        region = ET.SubElement(
            page_node,
            f"{{{NS}}}TextRegion",
            id=item.id,
            custom=f"readingOrder {{index:{item.reading_order};}} type:{item.region_type}",
        )
        ET.SubElement(region, f"{{{NS}}}Coords", points=_polygon_points(item.polygon))
        region_nodes[item.id] = region
    if not region_nodes:
        region = ET.SubElement(page_node, f"{{{NS}}}TextRegion", id="region_1")
        ET.SubElement(
            region,
            f"{{{NS}}}Coords",
            points=f"0,0 {page.width},0 {page.width},{page.height} 0,{page.height}",
        )
        region_nodes["region_1"] = region
    for line in page.lines:
        x1, y1, x2, y2 = line.bbox
        parent = region_nodes.get(line.region_id or "")
        if parent is None:
            parent = next(iter(region_nodes.values()))
        line_node = ET.SubElement(
            parent,
            f"{{{NS}}}TextLine",
            id=line.id,
            conf=f"{line.confidence:.5f}",
            custom=f"model:{line.model}; variant:{line.variant}",
        )
        polygon = line.polygon or [(x1, y1), (x2, y1), (x2, y2), (x1, y2)]
        ET.SubElement(line_node, f"{{{NS}}}Coords", points=_polygon_points(polygon))
        if line.baseline:
            ET.SubElement(
                line_node,
                f"{{{NS}}}Baseline",
                points=_polygon_points(line.baseline),
            )
        text_equiv = ET.SubElement(
            line_node, f"{{{NS}}}TextEquiv", index="0", conf=f"{line.confidence:.5f}"
        )
        ET.SubElement(text_equiv, f"{{{NS}}}Unicode").text = line.text
        for index, reading in enumerate(line.readings, start=1):
            if reading.text == line.text:
                continue
            alternative = ET.SubElement(
                line_node,
                f"{{{NS}}}TextEquiv",
                index=str(index),
                conf=f"{reading.confidence:.5f}",
                dataType=reading.kind.value,
                dataTypeDetails=reading.model,
            )
            ET.SubElement(alternative, f"{{{NS}}}Unicode").text = reading.text
    ET.ElementTree(root).write(destination, encoding="utf-8", xml_declaration=True)


ALTO_NS = "http://www.loc.gov/standards/alto/ns-v4#"


def _alto(page: PageResult, destination: Path, image_filename: str) -> None:
    root = ET.Element(f"{{{ALTO_NS}}}alto")
    description = ET.SubElement(root, f"{{{ALTO_NS}}}Description")
    source = ET.SubElement(description, f"{{{ALTO_NS}}}sourceImageInformation")
    ET.SubElement(source, f"{{{ALTO_NS}}}fileName").text = image_filename
    layout = ET.SubElement(root, f"{{{ALTO_NS}}}Layout")
    page_node = ET.SubElement(
        layout,
        f"{{{ALTO_NS}}}Page",
        ID=f"page_{page.page_index:04d}",
        WIDTH=str(page.width),
        HEIGHT=str(page.height),
        PHYSICAL_IMG_NR=str(page.page_index + 1),
    )
    print_space = ET.SubElement(
        page_node,
        f"{{{ALTO_NS}}}PrintSpace",
        HPOS="0",
        VPOS="0",
        WIDTH=str(page.width),
        HEIGHT=str(page.height),
    )
    block = ET.SubElement(print_space, f"{{{ALTO_NS}}}TextBlock", ID="block_1")
    for line in page.lines:
        x1, y1, x2, y2 = line.bbox
        line_node = ET.SubElement(
            block,
            f"{{{ALTO_NS}}}TextLine",
            ID=line.id,
            HPOS=str(x1),
            VPOS=str(y1),
            WIDTH=str(x2 - x1),
            HEIGHT=str(y2 - y1),
        )
        ET.SubElement(
            line_node,
            f"{{{ALTO_NS}}}String",
            ID=f"{line.id}_text",
            CONTENT=line.text,
            WC=f"{line.confidence:.5f}",
            HPOS=str(x1),
            VPOS=str(y1),
            WIDTH=str(x2 - x1),
            HEIGHT=str(y2 - y1),
        )
    ET.ElementTree(root).write(destination, encoding="utf-8", xml_declaration=True)


def _docx(result: DocumentResult, destination: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading("SchriftLotse", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(result.document.title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Jahr: {result.year or 'nicht angegeben'}")
    document.add_paragraph(f"Schrifthinweis: {result.script_hint.value}")
    document.add_heading("Originalgetreue Transkription", level=1)
    for page in result.pages:
        document.add_heading(f"Seite {page.page_index + 1}", level=2)
        for line in page.lines:
            paragraph = document.add_paragraph(line.text)
            if line.confidence < 0.9:
                paragraph.add_run(f"  [Sicherheit {line.confidence:.0%}]").italic = True
    document.add_page_break()
    document.add_heading("Lesefassung", level=1)
    for paragraph_text in reading_text(result).split("\n\n"):
        document.add_paragraph(paragraph_text)
    document.add_heading("Technischer Anhang", level=1)
    for page in result.pages:
        document.add_paragraph(
            f"Seite {page.page_index + 1}: {page.selected_model}, "
            f"Variante {page.selected_variant}, "
            f"mittlere Sicherheit {page.mean_confidence:.1%}, erwartete CER {page.expected_cer:.1%}"
        )
    document.save(str(destination))


def _pdf(result: DocumentResult, destination: Path) -> None:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="Cover", parent=styles["Title"], alignment=TA_CENTER, fontSize=26, leading=32
        )
    )
    styles.add(
        ParagraphStyle(
            name="Muted",
            parent=styles["BodyText"],
            textColor=colors.HexColor("#59636f"),
            fontSize=8,
        )
    )
    document = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Spacer(1, 45 * mm),
        Paragraph("SchriftLotse", styles["Cover"]),
        Spacer(1, 5 * mm),
        Paragraph(html.escape(result.document.title), styles["Title"]),
        Spacer(1, 8 * mm),
        Paragraph(
            f"Jahr: {result.year or 'nicht angegeben'} · Schrift: {result.script_hint.value}",
            styles["Normal"],
        ),
        PageBreak(),
        Paragraph("Originalgetreue Transkription", styles["Heading1"]),
    ]
    for page in result.pages:
        story.append(Paragraph(f"Seite {page.page_index + 1}", styles["Heading2"]))
        for line in page.lines:
            story.append(Paragraph(html.escape(line.text) or "&nbsp;", styles["BodyText"]))
        story.append(
            Paragraph(
                f"{html.escape(page.selected_model)} · {html.escape(page.selected_variant)} · "
                f"Sicherheit {page.mean_confidence:.0%}",
                styles["Muted"],
            )
        )
        story.append(Spacer(1, 4 * mm))
    story.extend([PageBreak(), Paragraph("Lesefassung", styles["Heading1"])])
    for paragraph_text in reading_text(result).split("\n\n"):
        story.append(
            Paragraph(html.escape(paragraph_text).replace("\n", "<br/>"), styles["BodyText"])
        )
        story.append(Spacer(1, 3 * mm))
    table_data = [["Seite", "Modell", "Variante", "Sicherheit", "erw. CER"]]
    for page in result.pages:
        table_data.append(
            [
                str(page.page_index + 1),
                page.selected_model,
                page.selected_variant,
                f"{page.mean_confidence:.1%}",
                f"{page.expected_cer:.1%}",
            ]
        )
    story.extend(
        [
            PageBreak(),
            Paragraph("Technischer Anhang", styles["Heading1"]),
            Table(table_data, repeatRows=1),
        ]
    )
    table = story[-1]
    if isinstance(table, Table):
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4d5c")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#c8d1d5")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
    document.build(story)


def export_document(result: DocumentResult, output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    result.output_dir = output_dir
    original = output_dir / "transkription_original.txt"
    readable = output_dir / "lesefassung.txt"
    json_path = output_dir / "result.json"
    docx_path = output_dir / "schriftlotse.docx"
    pdf_path = output_dir / "schriftlotse.pdf"
    original.write_text(original_text(result), encoding="utf-8")
    readable.write_text(reading_text(result), encoding="utf-8")
    json_path.write_text(
        json.dumps(result.model_dump(mode="json"), indent=2, ensure_ascii=False), encoding="utf-8"
    )
    _docx(result, docx_path)
    _pdf(result, pdf_path)
    pagexml_dir = output_dir / "pagexml"
    pagexml_dir.mkdir(exist_ok=True)
    alto_dir = output_dir / "alto"
    alto_dir.mkdir(exist_ok=True)
    images_dir = output_dir / "images"
    images_dir.mkdir(exist_ok=True)
    page_files: list[Path] = []
    alto_files: list[Path] = []
    image_files: list[Path] = []
    for page in result.pages:
        image_path = images_dir / f"seite_{page.page_index + 1:04d}.png"
        source_image = (
            page.prepared_path
            if page.prepared_path is not None and page.prepared_path.is_file()
            else page.source_path
        )
        if source_image.suffix.lower() == ".png":
            shutil.copy2(source_image, image_path)
        else:
            from schriftlotse.ingest import load_page

            load_page(source_image, page.source_page_index).save(image_path)
        page_path = pagexml_dir / f"seite_{page.page_index + 1:04d}.xml"
        relative_image = f"../images/{image_path.name}"
        _page_xml(result, page, page_path, relative_image)
        alto_path = alto_dir / f"seite_{page.page_index + 1:04d}.xml"
        _alto(page, alto_path, relative_image)
        page_files.append(page_path)
        alto_files.append(alto_path)
        image_files.append(image_path)
    escriptorium_archive = output_dir / "escriptorium-pagexml.zip"
    with zipfile.ZipFile(escriptorium_archive, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for path in [*page_files, *image_files]:
            zip_file.write(path, path.relative_to(output_dir))
    archive = output_dir / "schriftlotse-ergebnis.zip"
    files = [
        original,
        readable,
        json_path,
        docx_path,
        pdf_path,
        *page_files,
        *alto_files,
        *image_files,
        escriptorium_archive,
    ]
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for path in files:
            zip_file.write(path, path.relative_to(output_dir))
    return [*files, archive]


def export_search_results(rows: list[dict[str, object]], destination: Path) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    headers = ["Dokument", "Jahr", "Seite", "Text", "Trefferart", "Bewertung"]
    table_data = [headers] + [
        [
            str(row.get("document", "")),
            str(row.get("year", "")),
            str(row.get("page", "")),
            str(row.get("text", "")),
            str(row.get("reason", "")),
            str(row.get("score", "")),
        ]
        for row in rows
    ]
    pdf = SimpleDocTemplate(str(destination), pagesize=A4, rightMargin=10 * mm, leftMargin=10 * mm)
    table = Table(
        table_data, repeatRows=1, colWidths=[35 * mm, 15 * mm, 13 * mm, 75 * mm, 35 * mm, 18 * mm]
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4d5c")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    pdf.build([Paragraph("SchriftLotse – Suchergebnisse", getSampleStyleSheet()["Title"]), table])
    return destination
